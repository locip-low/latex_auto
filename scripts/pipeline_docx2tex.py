# -*- coding: utf-8 -*-
"""
Pipeline: Word (.docx) -> LaTeX (.tex) -> XeLaTeX -> PDF
Handles course design reports with no Word heading styles.
Usage: python pipeline_docx2tex.py <project_dir> [--docx name.docx] [--title Title]
"""
import sys, os, re, io, argparse, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    import docx
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

CLS_SOURCE = r"D:\BaiduNetdiskDownload\practice\bilibili\template\cumcmthesis.cls"

# ===== Helpers =====

def escape_tex(s):
    for ch, rep in [('&', r'\&'), ('%', r'\%'), ('$', r'\$'), ('#', r'\#'),
                     ('_', r'\_'), ('{', r'\{'), ('}', r'\}'),
                     ('~', r'\textasciitilde{}'), ('^', r'\textasciicircum{}')]:
        s = s.replace(ch, rep)
    return s

def detect_heading(text):
    m = re.match(r'^(\d+(?:\.\d+){0,2})\s{1,4}(.+)', text)
    if not m:
        return None, None
    num, title = m.group(1), m.group(2).strip()
    depth = num.count('.')
    return {0: 'section', 1: 'subsection', 2: 'subsubsection'}[depth], title

def convert_table_latex(table, caption, label):
    rows = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
    if not rows:
        return ''
    ncols = len(rows[0])
    colspec = '|' + 'c|' * ncols
    out = [
        r'\begin{table}[H]', r'  \centering',
        r'  \caption{' + caption + '}', r'  \label{' + label + '}',
        r'  \begin{tabular}{' + colspec + '}', r'    \hline',
    ]
    for row in rows:
        escaped = [escape_tex(c) if c else '---' for c in row]
        out.append(r'    ' + ' & '.join(escaped) + r' \\ \hline')
    out.append(r'  \end{tabular}')
    out.append(r'\end{table}')
    return '\n'.join(out)


def make_figure_placeholder(caption):
    """Generate a placeholder box for missing figures."""
    return (
        r'\begin{figure}[H]' + '\n'
        r'  \centering' + '\n'
        r'  \fbox{\begin{minipage}{0.8\textwidth}' + '\n'
        r'    \centering\vspace{3cm}' + '\n'
        r'    \small[' + caption + r' -- 图片待插入]' + '\n'
        r'    \vspace{3cm}' + '\n'
        r'  \end{minipage}}' + '\n'
        r'  \caption{' + caption + '}' + '\n'
        r'\end{figure}' + '\n'
    )


def fix_missing_figures(content, figures_dir):
    """Replace figure blocks with placeholders if the image file is missing."""
    existing = set()
    if os.path.isdir(figures_dir):
        for f in os.listdir(figures_dir):
            if f.endswith(('.png', '.jpg', '.jpeg', '.pdf')):
                existing.add(f)

    def replacer(m):
        block = m.group(0)
        fig_match = re.search(r'figures/(fig\d+)\.', block)
        if fig_match and (fig_match.group(1) + '.png') in existing:
            return block
        caption_match = re.search(r'\\caption\{(.+?)\}', block)
        caption = caption_match.group(1) if caption_match else '图片待插入'
        return make_figure_placeholder(caption)

    return re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', replacer, content, flags=re.DOTALL)


# ===== Pipeline =====

def pipeline(proj_dir, docx_name=None, title=None):
    # ---- Find docx ----
    if docx_name:
        docx_path = os.path.join(proj_dir, docx_name)
    else:
        docx_files = [f for f in os.listdir(proj_dir) if f.endswith('.docx')]
        if not docx_files:
            print(f"ERROR: No .docx found in {proj_dir}")
            return 1
        docx_path = os.path.join(proj_dir, docx_files[0])
        print(f"Step 1/5: Read docx -> {docx_files[0]}")

    doc = docx.Document(docx_path)
    paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    print(f"  Paragraphs: {len(paragraphs)}, Tables: {len(doc.tables)}")

    if title is None:
        title = paragraphs[0][1] if paragraphs else "Untitled"
    print(f"  Title: {title}")

    # ---- Detect TOC and body boundaries ----
    toc_idx = None
    for idx, (orig_i, text) in enumerate(paragraphs):
        if '目' in text and '录' in text and toc_idx is None:
            toc_idx = idx

    body_start = 0
    if toc_idx is not None:
        seen = set()
        for idx in range(toc_idx + 1, len(paragraphs)):
            level, t = detect_heading(paragraphs[idx][1])
            if level == 'section':
                if t in seen:
                    body_start = idx
                    break
                seen.add(t)
        if body_start == 0:
            body_start = toc_idx + 8

    code_start = None
    for idx, (orig_i, text) in enumerate(paragraphs):
        if re.match(r'7\.3\s+完整源代码', text) or re.match(r'附录.*源代码', text):
            if code_start is None:
                code_start = idx + 1

    print(f"Step 2/5: Parse structure -> body@{body_start}, code@{code_start}")

    # ---- Generate content.tex ----
    content_lines = []
    in_figure_section = False
    i = body_start
    while i < len(paragraphs):
        if code_start is not None and i >= code_start:
            break
        text = paragraphs[i][1]
        level, ht = detect_heading(text)

        if level:
            cmd = {'section': '\\section', 'subsection': '\\subsection',
                   'subsubsection': '\\subsubsection'}[level]
            content_lines.append(cmd + '{' + escape_tex(ht) + '}')
            in_figure_section = ('运行截图' in ht)
            i += 1
            continue

        if in_figure_section:
            fig_m = re.match(r'^(\d+)\.\s+(.+)', text)
            if fig_m and len(fig_m.group(2).strip()) <= 20:
                fig_num, fig_desc = fig_m.group(1), fig_m.group(2).strip()
                content_lines.append(r'\begin{figure}[H]')
                content_lines.append(r'  \centering')
                content_lines.append(r'  \includegraphics[width=0.85\textwidth]{figures/fig' + fig_num + '.png}')
                content_lines.append(r'  \caption{' + escape_tex(f'图{fig_num}: {fig_desc}界面') + '}')
                content_lines.append(r'  \label{fig:' + fig_num + '}')
                content_lines.append(r'\end{figure}')
                content_lines.append('')
                i += 1
                continue

        content_lines.append(escape_tex(text))
        content_lines.append('')
        i += 1

    content_text = '\n'.join(content_lines)
    content_text = re.sub(r'\n{3,}', '\n\n', content_text)
    content_path = os.path.join(proj_dir, 'content.tex')

    # Handle missing figures
    figures_dir = os.path.join(proj_dir, 'figures')
    content_text = fix_missing_figures(content_text, figures_dir)

    with open(content_path, 'w', encoding='utf-8') as f:
        f.write(content_text)
    print(f"Step 3/5: Write content.tex ({len(content_text)} chars)")

    # ---- Generate main.tex ----
    cls_dst = os.path.join(proj_dir, 'cumcmthesis.cls')
    if not os.path.exists(cls_dst) and os.path.exists(CLS_SOURCE):
        shutil.copy2(CLS_SOURCE, cls_dst)

    table_blocks = [convert_table_latex(t, f'表{ti+1}', f'tab:{ti+1}') for ti, t in enumerate(doc.tables)]
    tables_tex = '\n\n'.join(table_blocks)

    main_tex = r'''% !Mode:: "TeX:UTF-8"
% !TEX program  = xelatex
\documentclass[withoutpreface,bwprint]{cumcmthesis}

\usepackage{subfigure}
\usepackage{float}
\usepackage{url}

\usepackage{listings}
\usepackage{xcolor}
\definecolor{dkgreen}{rgb}{0,0.6,0}
\definecolor{gray}{rgb}{0.5,0.5,0.5}
\definecolor{mauve}{rgb}{0.58,0,0.82}
\lstset{
  frame=tb,
  aboveskip=3mm,
  belowskip=3mm,
  showstringspaces=false,
  columns=flexible,
  framerule=1pt,
  rulecolor=\color{gray!35},
  backgroundcolor=\color{gray!5},
  basicstyle={\small\ttfamily},
  numbers=left,
  numberstyle=\tiny\color{gray},
  keywordstyle=\color{blue},
  commentstyle=\color{dkgreen},
  stringstyle=\color{mauve},
  breaklines=true,
  breakatwhitespace=true,
  tabsize=3,
}

\bibliographystyle{plain}

\title{''' + title + r'''}

\begin{document}

\maketitle

\input{content.tex}

\newpage
\appendix

''' + tables_tex + r'''

\section{完整源代码}
\lstinputlisting[language=C, caption={''' + title + r'''完整源代码（C语言）}, label=code:full]{code/package_system.c}

\end{document}
'''

    main_path = os.path.join(proj_dir, 'main.tex')
    with open(main_path, 'w', encoding='utf-8') as f:
        f.write(main_tex)
    print(f"Step 4/5: Write main.tex")

    # ---- Extract code ----
    if code_start is not None:
        code_dir = os.path.join(proj_dir, 'code')
        os.makedirs(code_dir, exist_ok=True)
        code_text = '\n'.join(paragraphs[idx][1] for idx in range(code_start, len(paragraphs)))
        code_path = os.path.join(code_dir, 'package_system.c')
        with open(code_path, 'w', encoding='utf-8') as f:
            f.write(code_text)
        print(f"  Code extracted: code/package_system.c ({len(paragraphs) - code_start} lines)")

    # Ensure directories
    for sub in ['figures', 'code']:
        os.makedirs(os.path.join(proj_dir, sub), exist_ok=True)

    print(f"Step 5/5: Complete!")
    print(f"  -> Next: .\\scripts\\compile.ps1 -ProjectDir '{proj_dir}'")
    return 0


if __name__ == '__main__':
    ap = argparse.ArgumentParser(description='Word -> LaTeX pipeline for course reports')
    ap.add_argument('proj_dir', help='Project directory')
    ap.add_argument('--docx', help='Docx filename (auto-detect)')
    ap.add_argument('--title', help='Document title (auto-detect)')
    args = ap.parse_args()
    sys.exit(pipeline(args.proj_dir, args.docx, args.title))