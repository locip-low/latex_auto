# -*- coding: utf-8 -*-
"""Convert paper.docx to LaTeX — v11: fix keyword splitting by spaces."""
import docx, re, sys, os

DOCX_PATH = r"D:\vibecoding\latex_auto\test\_素材\paper.docx"
OUT_PATH = r"D:\vibecoding\latex_auto\test\content.tex"

doc = docx.Document(DOCX_PATH)

EXTRA_BOLD = []

def escape_tex_text(s):
    s = s.replace('&', r'\&'); s = s.replace('%', r'\%'); s = s.replace('$', r'\$')
    s = s.replace('#', r'\#'); s = s.replace('_', r'\_'); s = s.replace('{', r'\{')
    s = s.replace('}', r'\}'); s = s.replace('~', r'\textasciitilde{}'); s = s.replace('^', r'\textasciicircum{}')
    return s

CITE_MAP = {'1':'greenshields1935','2':'webster1958','3':'hochreiter1997','4':'hcm2016','5':'deb2002','6':'dijkstra1959'}

def process_citations(text):
    def repl(m):
        inner = m.group(1)
        nums = re.findall(r'\d+', inner)
        keys = [CITE_MAP.get(n, n) for n in nums]
        return r'\cite{' + ','.join(keys) + '}'
    return re.sub(r'\[(\s*[\d,\s]+)\]', repl, text)

def clean_normal_text(text):
    return escape_tex_text(process_citations(text))

def bold_keywords_in_text(text, keywords):
    """Wrap each keyword in \textbf{}."""
    for kw in sorted(keywords, key=len, reverse=True):
        if kw in text:
            text = text.replace(kw, r'\textbf{' + kw + '}')
    return text

def find_matching_paren(s, start):
    depth, i = 1, start
    while i < len(s) and depth > 0:
        if s[i] == '(': depth += 1
        elif s[i] == ')': depth -= 1
        i += 1
    return i - 1

def convert_sqrt(s):
    result, i = [], 0
    while i < len(s):
        if s[i:i+5] == 'sqrt(':
            end = find_matching_paren(s, i+5)
            result.append(r'\sqrt{' + s[i+5:end] + '}')
            i = end + 1
        else:
            result.append(s[i]); i += 1
    return ''.join(result)

def convert_sum(s):
    result, i = [], 0
    while i < len(s):
        if s[i:i+4] == 'sum(':
            end = find_matching_paren(s, i+4)
            result.append(r'\sum_{i=1}^{n}{' + s[i+4:end] + '}')
            i = end + 1
        else:
            result.append(s[i]); i += 1
    return ''.join(result)

def wrap_chinese_in_text(s):
    result, i = [], 0
    while i < len(s):
        ch = s[i]
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
            buf = ch; i += 1
            while i < len(s):
                ch2 = s[i]
                if '\u4e00' <= ch2 <= '\u9fff' or '\u3000' <= ch2 <= '\u303f' or '\uff00' <= ch2 <= '\uffef':
                    buf += ch2; i += 1
                else: break
            result.append(r'\text{' + buf + '}')
        else:
            result.append(ch); i += 1
    return ''.join(result)

def extract_formula_tex(text):
    t = text.strip()
    t = t.replace('[', r'\lbrack ').replace(']', r'\rbrack ')
    t = t.replace('·', r'\cdot ')
    t = wrap_chinese_in_text(t)
    subs = [
        ('v_f','v_{f}'),('k_j','k_{j}'),('q_m','q_{m}'),('g_i','g_{i}'),
        ('y_i','y_{i}'),('C_i','C_{i}'),('s_i','s_{i}'),('x_t','x_{t}'),
        ('h_t','h_{t}'),('w_ij','w_{ij}'),('k_ij','k_{ij}'),('g_ij','g_{ij}'),
        ('y_hat_i',r'\hat{y}_{i}'),('b_f','b_{f}'),('b_i','b_{i}'),
        ('b_o','b_{o}'),('W_f','W_{f}'),('W_i','W_{i}'),('W_o','W_{o}'),
        ('h_{t-1}','h_{t-1}'),
    ]
    for old, new in subs: t = t.replace(old, new)
    t = re.sub(r'_([a-zA-Z0-9]+)', r'_{\1}', t)
    t = t.replace('sigma(', r'\sigma(')
    t = convert_sqrt(t); t = convert_sum(t)
    t = t.replace('*', r' \times ').replace('<=', r' \leq ').replace('>=', r' \geq ').replace('...', r' \ldots ')
    for s in [r'\\times',r'\\leq',r'\\geq',r'\\ldots',r'\\sigma']:
        t = t.replace(s, s.replace('\\\\','\\'))
    return t

def convert_table_latex(table, idx):
    rows = [[cell.text.strip().replace('\n',' ') for cell in row.cells] for row in table.rows]
    if not rows: return ''
    colspec = 'c' * len(rows[0])
    out = [r'\begin{table}[H]', r'  \centering',
           r'  \caption{表' + str(idx) + ': ' + escape_tex_text(rows[0][0]) + '}',
           r'  \label{tab:' + str(idx) + '}',
           r'  \begin{tabular}{' + colspec + '}',
           r'    \toprule']
    for ri, row in enumerate(rows):
        escaped = [escape_tex_text(c) for c in row]
        out.append(r'    ' + ' & '.join(escaped) + r' \\')
        if ri == 0: out.append(r'    \midrule')
    out.append(r'    \bottomrule')
    out.append(r'  \end{tabular}')
    out.append(r'\end{table}')
    return '\n'.join(out)

def format_block(lines_list):
    result = '\n\n'.join(lines_list)
    for env in ['equation', 'figure', 'table', 'abstract']:
        result = re.sub(r'(\\begin\{' + env + r'\})\n\n+', r'\1\n', result)
        result = re.sub(r'\n\n+(\\end\{' + env + r'\})', r'\n\1', result)
        def fix_inner(m, e=env):
            inner = re.sub(r'\n\s*\n', '\n', m.group(2))
            return m.group(1) + inner + m.group(3)
        result = re.sub(r'(\\begin\{' + env + r'\})(.*?)(\\end\{' + env + r'\})',
                        fix_inner, result, flags=re.DOTALL)
    return result

# ===== Main loop =====
lines, table_idx = [], 0
title_skipped = False
in_abstract, keywords_done, in_refs = False, False, False
abstract_lines = []
keyword_list = []

i = 0
while i < len(doc.paragraphs):
    p = doc.paragraphs[i]
    text = p.text.strip()
    style = p.style.name
    if not text: i += 1; continue

    if not title_skipped and style == 'Normal' and i == 0:
        title_skipped = True; i += 1; continue

    if style == 'Heading 2' and '摘要' in text:
        in_abstract, keywords_done = True, False
        abstract_lines = []
        lines.append(r'\begin{abstract}'); i += 1; continue

    if in_abstract and not keywords_done:
        if '关键词' in text or '关键字' in text:
            kw_text = re.sub(r'关键[词字][：:]\s*', '', text)
            # Split by whitespace AND semicolons
            raw_kws = [k.strip() for k in re.split(r'[；;\s]+', kw_text) if k.strip()]
            keyword_list = raw_kws + EXTRA_BOLD

            # Bold keywords in abstract body lines
            for j in range(len(abstract_lines)):
                abstract_lines[j] = bold_keywords_in_text(abstract_lines[j], keyword_list)
            for al in abstract_lines:
                lines.append(al)

            # ALL keywords bold in keywords line
            kw_tex_items = [r'\textbf{' + escape_tex_text(kw) + '}' for kw in raw_kws]
            lines.append(r'\keywords{' + r'\qquad'.join(kw_tex_items) + '}')
            lines.append(r'\end{abstract}')
            keywords_done, in_abstract = True, False; i += 1; continue
        else:
            t = clean_normal_text(text)
            abstract_lines.append(t)
            i += 1; continue

    if style == 'Heading 1' and '参考' in text:
        in_refs = True
        lines.append(r'\nocite{*}')
        lines.append(r'\bibliography{ref}')
        i += 1; continue

    if in_refs: i += 1; continue

    fm = re.match(r'\u3010公式(\d+)\u3011\s*(.*)', text)
    if fm:
        eq_num, eq_body = fm.group(1), fm.group(2).strip()
        lines.append(r'\begin{equation}')
        lines.append('  ' + extract_formula_tex(eq_body))
        lines.append(r'  \label{eq:' + eq_num + '}')
        lines.append(r'\end{equation}')
        i += 1; continue

    figm = re.match(r'\[图\s*(\d+)\s+(.+?)\]', text)
    if figm:
        fig_num, caption = figm.group(1), clean_normal_text(figm.group(2).strip())
        lines.append(r'\begin{figure}[H]')
        lines.append(r'  \centering')
        lines.append(r'  \includegraphics[width=0.85\textwidth]{figures/fig' + fig_num + '.png}')
        lines.append(r'  \caption{' + caption + '}')
        lines.append(r'  \label{fig:' + fig_num + '}')
        lines.append(r'\end{figure}')
        i += 1; continue

    if style == 'Heading 1':
        h = re.sub(r'^[一二三四五六七八九十]+[、．.]?\s*', '', text)
        h = re.sub(r'^\d+(\.\d+)*\s*', '', h)
        lines.append(r'\section{' + escape_tex_text(h) + '}'); i += 1; continue

    if style == 'Heading 2':
        h = re.sub(r'^\d+(\.\d+)*\s*', '', text)
        if '摘要' in h: lines.append(r'\begin{abstract}'); i += 1; continue
        lines.append(r'\subsection{' + escape_tex_text(h) + '}'); i += 1; continue

    if style == 'Heading 3':
        h = re.sub(r'^\d+(\.\d+)*\s*', '', text)
        lines.append(r'\subsubsection{' + escape_tex_text(h) + '}'); i += 1; continue

    if style == 'Normal':
        t = clean_normal_text(text)
        table_ref = re.match(r'表\s*(\d+)\s+(.+)', t)
        if table_ref and table_idx < len(doc.tables):
            tb = doc.tables[table_idx]; table_idx += 1
            lines.append(convert_table_latex(tb, table_idx))
        else:
            lines.append(t)
        i += 1; continue

    i += 1

output = format_block(lines)
with open(OUT_PATH, 'w', encoding='utf-8') as f:
    f.write(output)

print(f'Written: {OUT_PATH}')
print(f'Keywords: {keyword_list}')